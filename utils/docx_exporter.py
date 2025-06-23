from docx import Document

def save_scene_as_docx(scene_text, filename):
    doc = Document()
    doc.add_heading("Generated Scene", 0)

    for line in scene_text.strip().split("\n"):
        doc.add_paragraph(line)

    doc.save(filename)
    return filename
